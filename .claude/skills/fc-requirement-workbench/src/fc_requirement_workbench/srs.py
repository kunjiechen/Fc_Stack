"""SRS structure generation and document rendering for Phase-3."""

from __future__ import annotations

from dataclasses import asdict, dataclass, field
from html import escape
from pathlib import Path
import re
from typing import Any

from .builder import EngineeringRequirement
from .rules import ValidationFinding
from .status import compute_requirement_status, render_status_label


SECTION_ORDER = [
    ("5.1 模式需求", "mode_state"),
    ("5.2 接口需求", "interface"),
    ("5.3 配置需求", "configuration"),
    ("5.4 诊断需求", "diagnostic"),
    ("6.1 时序需求", "timing"),
    ("6.2 安全等级需求", "safety"),
    ("6.3 编码规范需求", "coding"),
    ("6.4 资源消耗需求", "resource"),
]


@dataclass(frozen=True)
class SrsDocument:
    title: str
    module: str
    requirements: list[EngineeringRequirement]
    findings: list[ValidationFinding] = field(default_factory=list)
    overview: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "title": self.title,
            "module": self.module,
            "requirements": [req.to_dict() for req in self.requirements],
            "coverage_matrix": coverage_matrix(self.requirements),
            "trace_matrix": trace_matrix(self.requirements),
            "findings": [finding.to_dict() for finding in self.findings],
            "overview": self.overview,
        }


class SrsStructureGenerator:
    def build_document(
        self,
        requirements: list[EngineeringRequirement],
        module: str = "FC",
        findings: list[ValidationFinding] | None = None,
        overview: dict[str, Any] | None = None,
    ) -> SrsDocument:
        return SrsDocument(
            title=f"{module} 软件需求规范",
            module=module,
            requirements=sorted(requirements, key=lambda req: req.requirement_id),
            findings=findings or [],
            overview=overview or {},
        )

    def sections(self, document: SrsDocument) -> dict[str, list[EngineeringRequirement]]:
        result: dict[str, list[EngineeringRequirement]] = {key: [] for _, key in SECTION_ORDER}
        for req in document.requirements:
            key = _section_key(req)
            result.setdefault(key, []).append(req)
        return result


class MarkdownSrsRenderer:
    def render(self, document: SrsDocument) -> str:
        sections = SrsStructureGenerator().sections(document)
        safety_level = (document.overview or {}).get("safety_level", "QM")
        lines = _document_header_markdown(document)
        lines.extend(_purpose_markdown(document))
        lines.extend(_scope_markdown(document))
        lines.extend(_terms_markdown())
        lines.extend(_overview_markdown(document))
        lines.extend(["## 5 功能需求", ""])
        lines.append("本章描述模块必须实现的功能行为，包括模式、接口、配置和诊断。每条需求使用固定字段描述，以便后续生成设计、测试和追溯矩阵。")
        lines.append("")
        for heading, key in SECTION_ORDER[:4]:
            reqs = sections.get(key, [])
            lines.extend([f"### {heading}", ""])
            if not reqs:
                lines.append("无对应需求。")
                lines.append("")
                continue
            for req in reqs:
                lines.extend(_requirement_markdown(req, safety_level))

        lines.extend(["## 6 非功能需求", ""])
        for heading, key in SECTION_ORDER[4:]:
            reqs = sections.get(key, [])
            lines.extend([f"### {heading}", ""])
            if not reqs:
                lines.append("无对应需求。")
                lines.append("")
                continue
            for req in reqs:
                lines.extend(_requirement_markdown(req, safety_level))

        lines.extend(_risk_register_markdown(document))
        lines.extend(_sources_markdown(document.requirements))
        lines.extend(_requirement_list_markdown(document.requirements))
        lines.extend(_supporting_files_markdown())
        lines.extend(_review_and_release_guidance_markdown())
        return "\n".join(lines).rstrip() + "\n"


class HtmlSrsRenderer:
    def render(self, document: SrsDocument) -> str:
        markdown = MarkdownSrsRenderer().render(document)
        body = _markdown_subset_to_html(markdown)
        return "\n".join(
            [
                "<!doctype html>",
                "<html>",
                "<head>",
                "  <meta charset=\"utf-8\">",
                f"  <title>{escape(document.title)}</title>",
                "  <style>body{font-family:Arial,sans-serif;line-height:1.45;margin:40px;max-width:1080px} table{border-collapse:collapse;width:100%;margin:12px 0} th,td{border:1px solid #bbb;padding:6px;text-align:left;vertical-align:top} code{background:#f4f4f4;padding:1px 3px}</style>",
                "</head>",
                "<body>",
                body,
                "</body>",
                "</html>",
            ]
        )


class DocxSrsRenderer:
    def render_to_file(self, document: SrsDocument, path: str | Path) -> Path:
        from docx import Document
        from docx.shared import Inches

        output = Path(path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        doc.add_heading(document.title, level=0)
        _add_matrix_docx(
            doc,
            ["项目", "内容"],
            [
                ("文档编号", f"SRS-{_normalize_doc_token(document.module)}-001"),
                ("文档名称", document.title),
                ("模块名称", document.module),
                ("模块简称", document.module),
                ("文档状态", "Draft"),
                ("安全等级", (document.overview or {}).get("safety_level", "QM")),
            ],
        )
        doc.add_heading("1 目的", level=1)
        doc.add_paragraph(f"本文档定义 {document.module} 模块的软件需求。")
        doc.add_heading("2 适用范围", level=1)
        doc.add_paragraph(f"本文档适用于 {document.module} 模块的软件开发、评审、集成、测试和交付活动。")
        doc.add_heading("3 定义和缩写", level=1)
        _add_matrix_docx(doc, ["缩写", "英文全称", "中文说明"], _default_abbreviations())
        doc.add_heading("4 概述", level=1)
        doc.add_paragraph(f"{document.module} 模块需求由输入材料和需求语义对象生成。")

        sections = SrsStructureGenerator().sections(document)
        doc.add_heading("5 功能需求", level=1)
        for heading, key in SECTION_ORDER[:4]:
            doc.add_heading(heading, level=2)
            reqs = sections.get(key, [])
            if not reqs:
                doc.add_paragraph("本节暂无生成需求。")
                continue
            for req in reqs:
                _add_requirement_docx(doc, req)

        doc.add_heading("6 非功能需求", level=1)
        for heading, key in SECTION_ORDER[4:]:
            doc.add_heading(heading, level=2)
            reqs = sections.get(key, [])
            if not reqs:
                doc.add_paragraph("本节暂无生成需求。")
                continue
            for req in reqs:
                _add_requirement_docx(doc, req)

        doc.add_heading("7 风险与待确认问题", level=1)
        doc.add_heading("7.0 需求风险与待确认总表", level=2)
        _add_matrix_docx(
            doc,
            ["索引", "问题项", "问题/风险", "影响", "建议动作", "备注", "状态"],
            _risk_register_rows(document),
        )
        doc.add_heading("7.1 接口遗漏风险清单", level=2)
        _add_matrix_docx(
            doc,
            ["风险项", "风险等级", "说明", "建议动作"],
            _interface_omission_rows(document),
        )
        doc.add_heading("7.2 待确认接口清单", level=2)
        _add_matrix_docx(
            doc,
            ["接口名", "来源需求", "置信度", "待确认原因", "建议处理"],
            _pending_interface_rows(document),
        )
        doc.add_heading("7.3 不建议直接生成的低置信度接口", level=2)
        low_conf_rows = _low_confidence_interface_rows(document)
        if low_conf_rows:
            _add_matrix_docx(
                doc,
                ["接口名", "来源需求", "置信度", "不建议原因", "建议处理"],
                low_conf_rows,
            )
        else:
            doc.add_paragraph("本节为空——当前所有候选接口置信度均为中或高。")
        doc.add_heading("8 需求来源", level=1)
        _add_matrix_docx(doc, ["来源类别", "来源名称", "与本文档关系", "状态"], _source_rows(document.requirements))
        doc.add_heading("附录A 需求清单", level=1)
        _add_matrix_docx(doc, ["需求ID", "类别", "需求名称", "验证方式", "验证阶段", "状态"], _requirement_list_rows(document.requirements))
        doc.add_heading("附录B 支持和相关性文件", level=1)
        _add_matrix_docx(doc, ["序号", "文件名称", "文件编号/版本", "来源", "与本文档关系"], _supporting_file_rows())
        doc.add_heading("下一步：评审与发布引导", level=1)
        doc.add_paragraph("当需求状态为 Draft 时必须执行以下评审与发布引导：")
        for line in _review_and_release_guidance_lines():
            doc.add_paragraph(line, style="List Bullet")
        doc.save(output)
        return output


def coverage_matrix(requirements: list[EngineeringRequirement]) -> list[list[str]]:
    rows: list[list[str]] = []
    for req in requirements:
        rows.append(
            [
                req.requirement_id,
                _source_summary(req),
                "Warnings" if req.validation else "Passed",
                req.verification,
            ]
        )
    return rows


def trace_matrix(requirements: list[EngineeringRequirement]) -> list[list[str]]:
    rows: list[list[str]] = []
    for req in requirements:
        if not req.source:
            rows.append(["", req.requirement_id])
            continue
        for source in req.source:
            source_id = source.get("chunk_id") or source.get("document", "")
            rows.append([source_id, req.requirement_id])
    return rows


def _section_key(req: EngineeringRequirement) -> str:
    if req.requirement_type == "state":
        return "mode_state"
    if req.requirement_type == "functional":
        return "mode_state"
    if req.requirement_type == "timing":
        return "timing"
    return req.requirement_type


def _document_header_markdown(document: SrsDocument) -> list[str]:
    return [
        f"# 《{document.title}》",
        "",
        f"**{document.module}_软件需求规范**",
        "",
        f"**{document.module} Software Requirements Specification**",
        "",
        f"项目编号/Project number:{document.module}",
        "保密性/Security:**内部使用**",
        "",
        "**Document Properties**",
        "Status:**草稿**",
        "版本:**Draft**",
        "Author:待填写",
        "Created:待填写",
        "",
        "**Approved Versions**",
        "Current Document version **Draft** is **TBD**.",
        "",
        "**Approved Versions:**",
        "",
        "- TBD",
        "",
        "**Document Signatures**",
        "",
        "| 版本 | 状态 | 审批人 | 日期 | 意见 |",
        "| --- | --- | --- | --- | --- |",
        "| Draft | 草稿 | TBD | TBD | TBD |",
        "",
        "## 适用说明",
        "",
        f"本文档适用于 `{document.module}` 模块的软件需求定义。本文档仅描述软件应满足的需求，不描述详细设计方案、代码实现方案或测试用例步骤。",
        "",
        "---",
        "",
        "## 文档修订记录",
        "",
        "| 版本 | 日期 | 作者 | 变更说明 | 状态 |",
        "| --- | --- | --- | --- | --- |",
        "| Draft | 待填写 | 待填写 | 初版生成 | Draft |",
        "",
        "---",
        "",
        "## 目录",
        "",
        "- [1 目的](#1-目的)",
        "- [2 适用范围](#2-适用范围)",
        "- [3 定义和缩写](#3-定义和缩写)",
        "- [4 概述](#4-概述)",
        "- [5 功能需求](#5-功能需求)",
        "- [6 非功能需求](#6-非功能需求)",
        "- [7 风险与待确认问题](#7-风险与待确认问题)",
        "- [8 需求来源](#8-需求来源)",
        "- [附录A 需求清单](#附录a-需求清单)",
        "- [附录B 支持和相关性文件](#附录b-支持和相关性文件)",
        "- [下一步：评审与发布引导](#下一步评审与发布引导)",
        "",
        "---",
        "",
    ]


def _purpose_markdown(document: SrsDocument) -> list[str]:
    return [
        "## 1 目的",
        "",
        f"本文档定义 `{document.module}` 模块的软件需求，明确模块的功能边界、对外接口、状态行为、配置约束、诊断状态、时序要求、非功能约束和验证要求。",
        "",
        f"本文档作为 `{document.module}` 模块软件架构设计、详细设计、编码实现、单元测试、集成测试和系统测试的上游输入。所有正式需求均应具备需求 ID、来源、约束、验收准则和验证方式。",
        "",
        "---",
        "",
    ]


def _scope_markdown(document: SrsDocument) -> list[str]:
    return [
        "## 2 适用范围",
        "",
        f"本文档适用于 `{document.module}` 模块的软件开发、评审、集成、测试和交付活动。",
        "",
        "### 2.1 适用对象",
        "",
        "- 软件需求工程师",
        "- 软件架构和详细设计工程师",
        "- 软件开发工程师",
        "- 软件测试工程师",
        "- 功能安全工程师",
        "- 项目质量和配置管理人员",
        "",
        "### 2.2 适用范围",
        "",
        f"本文档覆盖 `{document.module}` 模块的软件功能、接口、配置、诊断、时序及相关非功能需求，并给出需求来源、验证方式、验证阶段和需求状态。本文档不展开详细设计方案、代码实现方案和测试用例步骤。",
        "",
        "---",
        "",
    ]


def _terms_markdown() -> list[str]:
    lines = [
        "## 3 定义和缩写",
        "",
        "### 3.1 定义",
        "",
        "| 术语 | 定义 |",
        "| --- | --- |",
        "| 对外支持行为 | 项目要求模块通过接口、配置或状态对外提供的软件行为。 |",
        "| 硬件能力 | 芯片或平台具备的能力，不自动等同于软件需求。 |",
        "| 软件责任 | 项目明确要求模块实现、拒绝、配置、验证或报告的行为。 |",
        "",
        "### 3.2 缩写",
        "",
        "| 缩写 | 英文全称 | 中文说明 |",
        "| --- | --- | --- |",
    ]
    lines.extend(f"| {abbr} | {full} | {cn} |" for abbr, full, cn in _default_abbreviations())
    lines.extend(["", "---", ""])
    return lines


def _overview_markdown(document: SrsDocument) -> list[str]:
    overview = document.overview or {}
    lines = [
        "## 4 概述",
        "",
        "本章仅保留理解需求所需的芯片和驱动背景信息，避免展开实现细节；正式软件责任以下文需求条目为准。",
        "",
    ]

    # ---- 4.1 外设芯片介绍 ----
    lines.extend(_chip_overview_markdown(document, overview))

    # ---- 4.2 驱动功能介绍 ----
    lines.extend(_driver_functions_markdown(document, overview))

    # ---- 4.3 外设引脚介绍 ----
    lines.extend(_pin_table_markdown(overview))

    # ---- 4.4 状态机介绍（按需生成） ----
    state_lines = _state_machine_markdown(document, overview)
    has_state_machine = bool(state_lines)
    if has_state_machine:
        lines.extend(state_lines)

    # ---- 通信参数（按需生成：仅 I2C/SPI 器件），编号跟随是否有状态机 ----
    comm_section_num = "4.5" if has_state_machine else "4.4"
    lines.extend(_communication_params_markdown(document, overview, section_num=comm_section_num))

    lines.extend(["---", ""])
    return lines


def _chip_overview_markdown(document: SrsDocument, overview: dict[str, Any]) -> list[str]:
    chip_intro = overview.get("chip_intro", "")
    chip_capabilities = overview.get("chip_capabilities") or overview.get("capability_items") or []
    lines = ["### 4.1 外设芯片介绍", ""]
    if chip_intro:
        lines.extend([chip_intro, ""])
    if chip_capabilities:
        lines.append("芯片支持以下功能：")
        lines.append("")
        lines.extend(f"- {item}" for item in chip_capabilities)
        lines.append("")
    if not chip_intro and not chip_capabilities:
        lines.extend([
            f"`{document.module}` 外设芯片用于扩展控制器外部控制和状态采集能力，"
            "驱动负责封装芯片访问、配置和运行时控制行为。",
            "",
        ])
    return lines


def _driver_functions_markdown(document: SrsDocument, overview: dict[str, Any]) -> list[str]:
    functions = overview.get("driver_functions") or [
        "初始化并配置外设工作参数。",
        "提供外设读写、状态处理和错误处理接口。",
    ]

    lines = [
        "### 4.2 驱动功能介绍",
        "",
        f"`{document.module}` 驱动应实现以下软件功能：",
        "",
    ]
    lines.extend(f"{idx}. {item}" for idx, item in enumerate(functions, start=1))
    lines.append("")
    return lines


def _pin_table_markdown(overview: dict[str, Any]) -> list[str]:
    pin_rows = overview.get("pin_rows") or []
    normalized: list[tuple[str, str, str]] = []
    for row in pin_rows:
        if len(row) >= 3:
            normalized.append((str(row[0]), str(row[1]), str(row[2])))
        else:
            normalized.append((str(row[0]), str(row[1]), ""))
    lines = [
        "### 4.3 外设引脚介绍",
        "",
        "| 引脚 | 方向 | Pin口功能 |",
        "| --- | --- | --- |",
    ]
    lines.extend(
        f"| {_escape_table_text(pin)} | {_escape_table_text(direction)} | {_escape_table_text(function)} |"
        for pin, direction, function in normalized
    )
    lines.append("")
    return lines


def _state_machine_markdown(document: SrsDocument, overview: dict[str, Any]) -> list[str]:
    sm_data = overview.get("state_machine")
    if not sm_data or not isinstance(sm_data, dict):
        return []
    states = sm_data.get("states", [])
    transitions = sm_data.get("transitions", [])
    diagram = sm_data.get("diagram", "")
    summary = sm_data.get("summary", "")
    if len(states) < 2 and len(transitions) < 2 and not diagram:
        return []
    lines = [
        "### 4.4 状态机介绍",
        "",
    ]
    if summary:
        lines.extend([summary, ""])
    if diagram:
        lines.extend([diagram, ""])
    if states:
        lines.extend([
            "| 状态 | 说明 | 进入条件 | 退出条件 |",
            "| --- | --- | --- | --- |",
        ])
        for state in states:
            lines.append(
                f"| {_escape_table_text(state.get('name', ''))} "
                f"| {_escape_table_text(state.get('description', ''))} "
                f"| {_escape_table_text(state.get('entry', ''))} "
                f"| {_escape_table_text(state.get('exit', ''))} |"
            )
        lines.append("")
    return lines


def _communication_params_markdown(document: SrsDocument, overview: dict[str, Any], section_num: str = "4.5") -> list[str]:
    comm = overview.get("communication")
    if not comm or not isinstance(comm, dict):
        return []
    bus_type = comm.get("bus_type", "")
    summary = comm.get("summary", "")
    speed_modes = comm.get("speed_modes", [])
    addressing = comm.get("device_addressing", "")
    timing_params = comm.get("timing_params", [])
    if not summary and not speed_modes and not timing_params and not addressing:
        return []
    lines = [
        f"### {section_num} {bus_type} 通信参数" if bus_type else f"### {section_num} 通信参数",
        "",
    ]
    if summary:
        lines.extend([summary, ""])
    if speed_modes:
        lines.append("关键通信参数：")
        lines.append("")
        lines.extend(f"- {mode}" for mode in speed_modes[:2])

    if addressing:
        lines.append(f"- 器件寻址：{addressing}")

    if timing_params:
        for param in timing_params[:4]:
            name = _escape_table_text(param.get("name", ""))
            minimum = _escape_table_text(param.get("min", ""))
            maximum = _escape_table_text(param.get("max", ""))
            unit = _escape_table_text(param.get("unit", ""))
            if minimum and maximum and maximum not in {"—", "-"}:
                value = f"{minimum}~{maximum} {unit}".strip()
            elif maximum and maximum not in {"—", "-"}:
                value = f"<= {maximum} {unit}".strip()
            else:
                value = f">= {minimum} {unit}".strip()
            lines.append(f"- {name}：{value}")
        lines.append("")

    return lines


def _requirement_markdown(req: EngineeringRequirement, safety_level: str = "QM") -> list[str]:
    # Single rendering contract for requirement items:
    # heading + status tags + prose description + type-specific bullet block.
    # Do not reintroduce per-item Markdown field tables here.
    category_tag = _short_category_tag(req.requirement_type)
    asil = safety_level
    method = _verification_method(req.verification)
    stage = _verification_stage(req.verification)
    verify_tag = method if method == stage else f"{method} / {stage}"
    status = _requirement_status(req)
    source = _source_summary(req)

    status_tags = [category_tag, asil, verify_tag, status]
    if source:
        status_tags.append(f"来源: {source}")
    status_bar = "`" + "` `".join(status_tags) + "`"

    lines = [
        f"#### {req.requirement_id} {req.title}",
        "",
        status_bar,
        "",
    ]

    if req.description:
        lines.extend([req.description, ""])

    # Constraint block — title & fields depend on requirement type
    block_title = _BLOCK_TITLES.get(req.requirement_type, "约束定义")
    block_fields = _BLOCK_FIELDS.get(req.requirement_type, [])

    description_text = (req.description or "").strip()
    bullets: list[str] = []
    for label, attr_name in block_fields:
        value = _block_attr_value(req, attr_name)
        if not value:
            continue
        # Skip when field value duplicates the prose description already shown above
        if value.strip() == description_text:
            continue
        bullets.append(f"- **{label}**：{_escape_inline(value)}")

    if bullets:
        lines.append(f"**{block_title}**")
        lines.append("")
        lines.extend(bullets)
        lines.append("")

    if req.validation:
        warnings = "；".join(item["message"] for item in req.validation)
        lines.append(f"> 评审提示：{_escape_inline(warnings)}")
        lines.append("")

    return lines


# ---------------------------------------------------------------------------
# New format: block title & field mappings per requirement type
# ---------------------------------------------------------------------------

_BLOCK_TITLES: dict[str, str] = {
    "interface": "接口约束",
    "functional": "功能约束",
    "configuration": "配置约束",
    "diagnostic": "诊断约束",
    "state": "状态约束",
    "timing": "时序约束",
    "safety": "约束定义",
    "coding": "约束定义",
    "resource": "约束定义",
}

_BLOCK_FIELDS: dict[str, list[tuple[str, str]]] = {
    "interface": [
        ("前置条件", "pre_condition"),
        ("触发条件", "trigger"),
        ("输入", "input"),
        ("输出", "output"),
        ("异常处理", "exception"),
        ("验收准则", "verification"),
    ],
    "functional": [
        ("输入", "input"),
        ("输出", "output"),
        ("行为边界", "constraint"),
        ("异常处理", "exception"),
        ("验收准则", "verification"),
    ],
    "configuration": [
        ("配置项", "constraint"),
        ("前置条件", "pre_condition"),
        ("验收准则", "verification"),
    ],
    "diagnostic": [
        ("触发条件", "trigger"),
        ("输入", "input"),
        ("输出", "output"),
        ("异常处理", "exception"),
        ("行为边界", "constraint"),
        ("验收准则", "verification"),
    ],
    "state": [
        ("状态描述", "description"),
        ("触发条件", "trigger"),
        ("转换规则", "constraint"),
        ("验收准则", "verification"),
    ],
    "timing": [
        ("时序约束", "constraint"),
        ("验收准则", "verification"),
    ],
    "safety": [
        ("约束内容", "description"),
        ("适用范围", "constraint"),
        ("验收准则", "verification"),
    ],
    "coding": [
        ("约束内容", "description"),
        ("适用范围", "constraint"),
        ("验收准则", "verification"),
    ],
    "resource": [
        ("约束内容", "description"),
        ("适用范围", "constraint"),
        ("验收准则", "verification"),
    ],
}


def _short_category_tag(req_type: str) -> str:
    return {
        "functional": "功能需求",
        "state": "状态需求",
        "interface": "接口需求",
        "configuration": "配置需求",
        "diagnostic": "诊断需求",
        "timing": "时序需求",
        "safety": "安全需求",
        "coding": "编码需求",
        "resource": "资源需求",
    }.get(req_type, req_type)


def _block_attr_value(req: EngineeringRequirement, attr_name: str) -> str:
    if attr_name == "verification":
        return _acceptance_criteria(req)
    return (getattr(req, attr_name, "") or "").strip()


def _escape_inline(value: str) -> str:
    return str(value).replace("|", "\\|").replace("\n", "；")


def _sources_markdown(requirements: list[EngineeringRequirement]) -> list[str]:
    return [
        "## 8 需求来源",
        "",
        "| 来源类别 | 来源名称 | 与本文档关系 | 状态 |",
        "| --- | --- | --- | --- |",
        *_source_rows_markdown(requirements),
        "",
        "---",
        "",
    ]


def _requirement_list_markdown(requirements: list[EngineeringRequirement]) -> list[str]:
    lines = [
        "## 附录A 需求清单",
        "",
        "| 需求ID | 类别 | 需求名称 | 验证方式 | 验证阶段 | 状态 |",
        "| --- | --- | --- | --- | --- | --- |",
    ]
    lines.extend(
        f"| {req.requirement_id} | {_category_label(req)} | {_escape_table_text(req.title)} | {_verification_method(req.verification)} | {_verification_stage(req.verification)} | {_requirement_status(req)} |"
        for req in requirements
    )
    lines.extend(["", "---", ""])
    return lines


def _risk_register_markdown(document: SrsDocument) -> list[str]:
    risk_rows = _risk_register_rows(document)
    interface_risk_rows = _interface_omission_rows(document)
    pending_interface_rows = _pending_interface_rows(document)
    low_conf_rows = _low_confidence_interface_rows(document)
    lines = [
        "## 7 风险与待确认问题",
        "",
        "本章汇总当前需求版本中仍需项目确认、补料或后续评审关闭的事项，结构和评审方式与架构阶段保持一致，便于后续继承评审结论。",
        "",
        "### 7.0 需求风险与待确认总表",
        "",
        "| 索引 | 问题项 | 问题/风险 | 影响 | 建议动作 | 备注 | 状态 |",
        "| --- | --- | --- | --- | --- | --- | --- |",
    ]
    for row in risk_rows:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} | {row[4]} | {row[5]} | {row[6]} |")
    lines.extend([
        "",
        "### 7.1 接口遗漏风险清单",
        "",
        "| 风险项 | 风险等级 | 说明 | 建议动作 |",
        "| --- | --- | --- | --- |",
    ])
    for row in interface_risk_rows:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} |")
    lines.extend([
        "",
        "### 7.2 待确认接口清单",
        "",
        "| 接口名 | 来源需求 | 置信度 | 待确认原因 | 建议处理 |",
        "| --- | --- | --- | --- | --- |",
    ])
    for row in pending_interface_rows:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} | {row[4]} |")
    lines.extend([
        "",
        "### 7.3 不建议直接生成的低置信度接口",
        "",
    ])
    if low_conf_rows:
        lines.extend([
            "| 接口名 | 来源需求 | 置信度 | 不建议原因 | 建议处理 |",
            "| --- | --- | --- | --- | --- |",
        ])
        for row in low_conf_rows:
            lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} | {row[4]} |")
    else:
        lines.append("本节为空——当前所有候选接口置信度均为中或高。")
    lines.extend(["", "---", ""])
    return lines


def _risk_register_rows(document: SrsDocument) -> list[tuple[str, str, str, str, str, str, str]]:
    rows: list[tuple[str, str, str, str, str, str, str]] = []
    seen: set[tuple[str, str, str]] = set()
    index = 1

    for req in document.requirements:
        status = compute_requirement_status(req)
        if status == "ready":
            continue
        issue_type, summary, action = _summarize_requirement_pending(req, status)
        affected = req.requirement_id
        dedupe_key = (issue_type, summary, affected)
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        rows.append(
            (
                f"R{index}",
                _escape_table_text(issue_type),
                _escape_table_text(summary),
                _escape_table_text(affected),
                _escape_table_text(action),
                "",
                "待评审",
            )
        )
        index += 1

    for finding in document.findings:
        if finding.status != "failed" or finding.requirement_ids:
            continue
        issue_type = _pending_type_label_from_finding(finding.rule_group)
        summary = finding.message.strip() or f"{finding.rule_group} 存在待确认问题"
        action = _pending_action_from_finding(finding.rule_group)
        dedupe_key = (issue_type, summary, "-")
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        rows.append(
            (
                f"R{index}",
                _escape_table_text(issue_type),
                _escape_table_text(summary),
                "-",
                _escape_table_text(action),
                "",
                "待评审",
            )
        )
        index += 1

    rows = rows[:20]
    rows.append(("R-OTHER", "其他", "用户补充的其他建议或风险。", "用户填写。", "用户填写。", "无其他建议。", "待评审"))
    return rows


def _interface_omission_rows(document: SrsDocument) -> list[tuple[str, str, str, str]]:
    rows: list[tuple[str, str, str, str]] = []
    for req in document.requirements:
        if req.requirement_type != "interface":
            continue
        if compute_requirement_status(req) == "ready" and not req.validation:
            continue
        rows.append(
            (
                _escape_table_text(req.title),
                "中",
                _escape_table_text(f"{req.requirement_id} 当前仍存在接口边界、命名、依赖或责任归属待确认项。"),
                "结合风险总表结论确认后保留、修改或移除该接口。",
            )
        )
    if not rows:
        rows.append(("无", "低", "当前无新增接口遗漏风险。", "无"))
    return rows


def _pending_interface_rows(document: SrsDocument) -> list[tuple[str, str, str, str, str]]:
    rows: list[tuple[str, str, str, str, str]] = []
    for req in document.requirements:
        if req.requirement_type != "interface":
            continue
        status = compute_requirement_status(req)
        if status == "ready" and not req.validation:
            continue
        _, summary, action = _summarize_requirement_pending(req, status)
        confidence = "中" if status != "ready" else "高"
        rows.append(
            (
                _escape_table_text(req.title),
                _escape_table_text(req.requirement_id),
                confidence,
                _escape_table_text(summary),
                _escape_table_text(action),
            )
        )
    if not rows:
        rows.append(("无", "无", "高", "当前无待确认接口。", "无"))
    return rows


def _low_confidence_interface_rows(document: SrsDocument) -> list[tuple[str, str, str, str, str]]:
    rows: list[tuple[str, str, str, str, str]] = []
    for req in document.requirements:
        if req.requirement_type != "interface":
            continue
        if compute_requirement_status(req) != "open_issue":
            continue
        rows.append(
            (
                _escape_table_text(req.title),
                _escape_table_text(req.requirement_id),
                "低",
                _escape_table_text(f"{req.title} 当前仍存在较大的责任边界或接口契约不确定性。"),
                "建议待项目输入收敛后再正式生成该接口。",
            )
        )
    return rows


def _review_and_release_guidance_lines() -> list[str]:
    return [
        "推荐评审方式 1：直接修改上方风险表中的`状态`和`备注`。",
        "推荐评审方式 2：在当前窗口回复，例如`R1、R3 已评审；R5 待修改，备注：接口名统一为 xxx`。",
        "如果所有风险项均认可，可回复：`全部已评审，R-OTHER 无其他建议，直接发布`。",
        "如果某项需要修改，可回复：`R5 待修改，备注：xxx`。",
        "修改完成后仍保持当前版本的`Draft`，直到所有真实风险项均为`已评审`后发布为`Released`。",
        "草稿评审发布不升级版本；只有正式需求文件 + 新架构/下游交付基线发布时才升级到下一版本。",
    ]


def _review_and_release_guidance_markdown() -> list[str]:
    lines = ["## 下一步：评审与发布引导", "", "当需求状态为 `Draft` 时必须执行以下评审与发布引导：", ""]
    lines.extend(f"- {item}" for item in _review_and_release_guidance_lines())
    lines.extend(["", "---", ""])
    return lines


def _supporting_files_markdown() -> list[str]:
    lines = [
        "## 附录B 支持和相关性文件",
        "",
        "| 序号 | 文件名称 | 文件编号/版本 | 来源 | 与本文档关系 |",
        "| --- | --- | --- | --- | --- |",
    ]
    lines.extend(
        f"| {seq} | {name} | {version} | {source} | {relation} |"
        for seq, name, version, source, relation in _supporting_file_rows()
    )
    lines.append("")
    return lines


def _category_label(req: EngineeringRequirement) -> str:
    return _short_category_tag(req.requirement_type)


def _acceptance_criteria(req: EngineeringRequirement) -> str:
    if req.verification:
        return req.verification.rstrip(".。")
    return "通过评审、分析或测试确认需求行为满足。"


def _verification_method(verification: str) -> str:
    lowered = verification.lower()
    methods: list[str] = []
    if "review" in lowered or "inspection" in lowered or "评审" in verification or "检查" in verification:
        methods.append("Review")
    if "analysis" in lowered or "分析" in verification:
        methods.append("Analysis")
    if "test" in lowered or "测试" in verification:
        methods.append("Test")
    if not methods:
        methods.append("Review")
    return "/".join(dict.fromkeys(methods))


def _verification_stage(verification: str) -> str:
    lowered = verification.lower()
    stages: list[str] = []
    if "unit" in lowered or "ut" in lowered or "功能测试" in verification or "边界测试" in verification:
        stages.append("UT")
    if "integration" in lowered or "it" in lowered or "interface" in lowered or "接口" in verification or "集成" in verification:
        stages.append("IT")
    if "system" in lowered or "st" in lowered or "系统" in verification:
        stages.append("ST")
    if "review" in lowered or "inspection" in lowered or "analysis" in lowered or "评审" in verification or "分析" in verification:
        stages.append("Review")
    if not stages:
        stages.append("UT")
    return "/".join(dict.fromkeys(stages))


def _requirement_status(req: EngineeringRequirement) -> str:
    return render_status_label(compute_requirement_status(req))


def _summarize_requirement_pending(
    req: EngineeringRequirement,
    status: str,
) -> tuple[str, str, str]:
    text_blob = " ".join(
        [
            req.description,
            req.constraint,
            req.pre_condition,
            req.trigger,
            req.input,
            req.output,
            req.exception,
            req.verification,
        ]
    )
    if not req.source:
        return (
            "来源待补充",
            f"{req.title} 缺少可追溯来源，当前不能作为稳定输入下传。",
            "补充 datasheet、项目需求或追溯依据，并重新评审该需求。",
        )
    if any(kw in text_blob for kw in ("项目安全计划定义", "项目开发计划定义", "项目资源计划定义")):
        return (
            "项目输入待确认",
            f"{req.title} 仍依赖项目侧输入确认，边界尚未完全收敛。",
            "明确默认值、范围、命名、所有权或模式支持后更新需求。",
        )
    if req.validation:
        first_warning = str(req.validation[0].get("message", "")).strip() or f"{req.title} 存在规则校验遗留项。"
        return (
            "规则校验风险",
            first_warning,
            "根据评审提示修改需求表述，并确认验证方式和约束字段完整。",
        )
    if status == "open_issue":
        return (
            "行为边界待确认",
            f"{req.title} 仍存在未闭合的行为边界或责任归属问题。",
            "在需求评审中确认责任边界，并同步更新下游设计输入。",
        )
    return (
        "Draft 待收敛",
        f"{req.title} 当前仍为 Draft，建议在进入架构前继续收敛。",
        "结合补料结果完善需求，使其达到可稳定下传的状态。",
    )


def _pending_type_label_from_finding(rule_group: str) -> str:
    mapping = {
        "trace": "来源覆盖风险",
        "ownership": "所有权待确认",
        "configuration": "配置待确认",
        "dependency": "依赖待确认",
        "consistency": "一致性风险",
        "constraint": "约束冲突风险",
        "naming": "命名待收敛",
        "completeness": "字段缺失风险",
    }
    return mapping.get(rule_group, "全局待确认事项")


def _pending_action_from_finding(rule_group: str) -> str:
    mapping = {
        "trace": "补充来源或明确追溯关系后再下传。",
        "ownership": "补充接口、引脚或责任归属定义。",
        "configuration": "补充项目配置值、默认值或范围。",
        "dependency": "补充依赖接口、外部条件或协作边界。",
        "consistency": "统一相互冲突的需求表述和约束。",
        "constraint": "确认冲突约束的优先级和生效版本。",
        "naming": "统一命名后再进入下游设计。",
        "completeness": "补齐缺失字段并复核可验证性。",
    }
    return mapping.get(rule_group, "在需求评审中确认后关闭。")


def _escape_table_text(value: str) -> str:
    return str(value).replace("|", "\\|").replace("\n", "<br>")


def _normalize_doc_token(value: str) -> str:
    token = "".join(ch for ch in value.upper() if ch.isalnum())
    return token or "FC"


def _default_abbreviations() -> list[tuple[str, str, str]]:
    return [
        ("SRS", "Software Requirement Specification", "软件需求规范"),
        ("SDD", "Software Design Description", "软件设计说明"),
        ("UT", "Unit Test", "单元测试"),
        ("IT", "Integration Test", "集成测试"),
        ("ST", "System Test", "系统测试"),
        ("QM", "Quality Management", "质量管理等级"),
        ("GPIO", "General Purpose Input/Output", "通用输入输出"),
        ("I2C", "Inter-Integrated Circuit", "双线串行总线"),
    ]


def _source_rows_markdown(requirements: list[EngineeringRequirement]) -> list[str]:
    rows = _source_rows(requirements)
    return [
        f"| {_escape_table_text(category)} | {_escape_table_text(name)} | {_escape_table_text(relation)} | {_escape_table_text(status)} |"
        for category, name, relation, status in rows
    ]


def _source_rows(requirements: list[EngineeringRequirement]) -> list[tuple[str, str, str, str]]:
    seen: set[str] = set()
    rows: list[tuple[str, str, str, str]] = []
    for req in requirements:
        for source in req.source:
            name = _public_source_name(source)
            if name in seen:
                continue
            seen.add(name)
            rows.append(("输入材料", name, "需求来源", "已接入"))
    if not rows:
        rows.append(("输入材料", "待补充", "需求来源", "待补充"))
    return rows


def _requirement_list_rows(
    requirements: list[EngineeringRequirement],
) -> list[tuple[str, str, str, str, str, str]]:
    return [
        (
            req.requirement_id,
            _category_label(req),
            req.title,
            _verification_method(req.verification),
            _verification_stage(req.verification),
            _requirement_status(req),
        )
        for req in requirements
    ]


def _supporting_file_rows() -> list[tuple[str, str, str, str, str]]:
    return [
        ("1", "项目需求文档", "待填写", "项目输入", "软件需求来源"),
        ("2", "Datasheet", "待填写", "芯片资料", "芯片能力、引脚、状态和时序约束来源"),
        ("3", "项目开发规范", "待填写", "项目规范", "编码、资源和过程约束来源"),
    ]


def ensure_default_engineering_requirements(
    requirements: list[EngineeringRequirement],
    module: str,
    safety_level: str = "QM",
    *,
    mainfunction_required: bool = True,
) -> tuple[list[EngineeringRequirement], list[ValidationFinding]]:
    result = _with_default_diagnostic_requirements(requirements, module)
    token = _normalize_doc_token(module)

    # Ensure lifecycle interfaces (Init + MainFunction) for IoExtDev drivers.
    # Respect the profile's mainfunction_required flag — some driver types
    # (e.g. CAN/LIN transceivers) do not need a periodic MainFunction.
    result = _with_default_lifecycle_requirements(result, module, token, mainfunction_required=mainfunction_required)

    existing = {req.requirement_type for req in result}
    if "safety" not in existing:
        result.append(
            EngineeringRequirement(
                requirement_id=f"SRS-{token}-SAFE-0001",
                semantic_id=f"DEFAULT-{token}-SAFE-0001",
                requirement_type="safety",
                title="功能安全等级要求",
                description=f"软件需求按 {safety_level} 等级管理。所有软件接口应具备参数有效性检查与开发错误检测能力；与安全目标相关的功能应满足 {safety_level} 的诊断覆盖要求；故障处理与恢复策略应符合 {safety_level} 等级的安全状态定义。",
                constraint=f"{safety_level} 等级的 ASIL 分解、诊断覆盖目标、安全状态定义和验证策略由项目安全计划定义。",
                verification="通过安全需求评审、故障注入测试和安全分析验证。",
                source=[{"document": "Project Input", "chunk_id": "SAFETY-LEVEL", "evidence": f"Safety level specified as {safety_level}."}],
            )
        )
    if "coding" not in existing:
        result.append(
            EngineeringRequirement(
                requirement_id=f"SRS-{token}-CODE-0001",
                semantic_id=f"DEFAULT-{token}-CODE-0001",
                requirement_type="coding",
                title="编码规范符合性要求",
                description=f"Gp_{_short_module(module)} 驱动编码应符合《FC 开发指南(C语言)》规范要求。",
                constraint="项目编码规范版本、MISRA/静态检查规则、复杂度阈值由项目开发计划定义。",
                verification="通过代码评审、静态分析和编码规范检查验证。",
                source=[{"document": "SRS Template", "chunk_id": "DEFAULT-CODING-STANDARD", "evidence": "Default coding standard requirement generated from template."}],
            )
        )
    if "resource" not in existing:
        result.append(
            EngineeringRequirement(
                requirement_id=f"SRS-{token}-RES-0001",
                semantic_id=f"DEFAULT-{token}-RES-0001",
                requirement_type="resource",
                title="资源消耗要求",
                description=f"Gp_{_short_module(module)} 驱动应满足 ROM 消耗小于5KB, RAM 消耗小于2KB。",
                constraint="ROM/RAM/栈/CPU 预算由项目资源计划定义，编译后通过 map 文件统计验证。",
                verification="通过编译后资源统计（map 文件分析）和集成测试验证。",
                source=[{"document": "SRS Template", "chunk_id": "DEFAULT-RESOURCE-BUDGET", "evidence": "Default resource budget requirement — project must fill actual limits."}],
            )
        )

    # ---- Interface dedup: remove duplicate function names, keep most complete ----
    result, dedup_findings = dedup_interface_requirements(result, module)
    return result, dedup_findings


def _with_default_lifecycle_requirements(
    requirements: list[EngineeringRequirement],
    module: str,
    token: str,
    *,
    mainfunction_required: bool = True,
) -> list[EngineeringRequirement]:
    """Ensure Init and MainFunction lifecycle interfaces exist for IoExtDev drivers.

    Respects ``mainfunction_required`` — when False (e.g. CAN/LIN transceivers),
    MainFunction is NOT injected even if absent.
    """
    result = list(requirements)
    names = " ".join(
        f"{req.title} {req.function_name} {req.description}".lower()
        for req in requirements
    )

    # Only match Init as an end-of-name token, not Gp_ prefix in unrelated names
    has_init = any(kw in names for kw in ("_init", "驱动初始化", "初始化接口"))
    has_init = has_init or any(
        req.function_name.endswith("Init") or req.function_name.endswith("_Init")
        for req in requirements
    )

    if not has_init:
        result.append(
            EngineeringRequirement(
                requirement_id=f"SRS-{token}-IF-9002",
                semantic_id=f"DEFAULT-{token}-LIFECYCLE-INIT",
                requirement_type="interface",
                title="驱动初始化",
                description="软件应提供初始化接口，用于加载项目配置，建立运行时上下文，执行芯片上电序列（包括模式锁存、外部元件稳定等待），并在配置非法或初始化失败时返回定义错误。",
                pre_condition="芯片未初始化或处于复位/睡眠状态。",
                output="初始化后的运行时上下文或错误码。",
                exception="配置无效、底层访问失败、时序等待超时时返回对应错误，并保持芯片处于安全状态。",
                constraint="初始化接口必须可重入；重复调用应在不破坏已有状态的前提下返回已初始化状态。",
                verification="通过初始化测试验证：上电后调用初始化接口应返回成功，芯片应进入活动模式；非法配置应返回错误。",
                function_name=f"Gp_{_short_module(module)}_Init",
                source=[{"document": "SRS Template", "chunk_id": "DEFAULT-LIFECYCLE-INIT", "evidence": "Init is a mandatory lifecycle interface for every IoExtDev driver."}],
            )
        )

    has_mainfunction = any(
        req.function_name.endswith("MainFunction") or req.function_name.endswith("_MainFunction")
        for req in requirements
    )
    if not has_mainfunction and mainfunction_required:
        result.append(
            EngineeringRequirement(
                requirement_id=f"SRS-{token}-IF-9003",
                semantic_id=f"DEFAULT-{token}-LIFECYCLE-MAINFUNCTION",
                requirement_type="interface",
                title="周期主函数",
                description="软件应提供 MainFunction 接口，用于周期推进运行时状态、刷新输出信号、采集诊断输入和更新故障状态。接口不得执行长时间阻塞操作。",
                pre_condition="驱动已完成初始化。",
                trigger="由上层周期调度器按固定周期调用。",
                output="更新后的运行时状态、诊断状态和输出刷新结果。",
                exception="底层访问失败时记录诊断事件并保持上次有效状态，不阻塞调度。",
                constraint="MainFunction 必须满足项目实时性约束，单次调用最大耗时不应超过项目定义的调度周期预算。",
                verification="通过周期调度测试验证：在定义周期内调用 MainFunction，应完成输入采集、状态推进和输出刷新。",
                function_name=f"Gp_{_short_module(module)}_MainFunction",
                source=[{"document": "SRS Template", "chunk_id": "DEFAULT-LIFECYCLE-MAINFUNCTION", "evidence": "MainFunction is a mandatory lifecycle interface for every IoExtDev driver."}],
            )
        )

    return result


def _short_module(module: str) -> str:
    """Strip Gp_ prefix if present for consistent ID generation."""
    name = module.replace("Gp_", "").replace("GP_", "")
    return name or module


def _with_default_diagnostic_requirements(
    requirements: list[EngineeringRequirement],
    module: str,
) -> list[EngineeringRequirement]:
    result = list(requirements)
    token = _normalize_doc_token(module)
    names = " ".join(
        f"{req.title} {req.description} {req.constraint} {req.exception} {req.input} {req.output}".lower()
        for req in requirements
    )

    # DET check: only skip if a *dedicated* DET requirement exists, not just mentions
    has_det_requirement = any(
        "开发错误检测" in (getattr(req, "title", "") + getattr(req, "description", ""))
        for req in requirements
    )
    if not has_det_requirement:
        result.append(
            EngineeringRequirement(
                requirement_id=f"SRS-{token}-DIAG-9001",
                semantic_id=f"DEFAULT-{token}-DIAG-DET",
                requirement_type="diagnostic",
                title="开发错误检测",
                description="软件应提供开发错误检测能力，对未初始化访问、空指针、非法参数和非法调用顺序进行检测，并按项目约定进行 DET 上报或错误返回。",
                exception="未初始化访问、空指针、非法参数、非法状态调用时，应拒绝继续执行当前请求并保持无关状态不被破坏。",
                constraint="DET 为必备需求；若项目未集成 DET 模块，也应保留等效的开发错误检测和错误返回语义。",
                verification="通过接口测试、边界测试和故障注入验证：触发未初始化访问、空指针、非法参数和非法调用顺序时，应产生定义的 DET 上报或错误返回，且内部状态保持一致。",
                source=[{"document": "SRS Template", "chunk_id": "DEFAULT-DIAG-DET", "evidence": "DET is a mandatory baseline diagnostic requirement."}],
            )
        )

    has_fault_like_behavior = any(
        keyword in names for keyword in ("fault", "diag", "diagnostic", "error", "interrupt", "中断", "故障", "诊断")
    )
    has_fault_read = any(
        keyword in names for keyword in ("getdevfault", "getdiag", "故障读取", "故障信息读取", "诊断读取", "故障状态读取", "诊断状态读取")
    )
    if has_fault_like_behavior and not has_fault_read:
        result.append(
            EngineeringRequirement(
                requirement_id=f"SRS-{token}-IF-9001",
                semantic_id=f"DEFAULT-{token}-DIAG-READ",
                requirement_type="interface",
                title="故障诊断信息读取接口",
                description="软件应提供故障或诊断状态读取接口，用于向上层返回当前故障状态、诊断结果或中断相关状态信息。",
                output="故障位、诊断状态字或项目定义的状态结果。",
                exception="当驱动未初始化、参数非法或底层状态不可获取时，应返回定义错误，并保留最近一次有效诊断状态或给出无效标识。",
                constraint="当模块存在故障检测、诊断判定或中断异常处理时，应提供对应的故障/诊断读取接口或等效状态读取机制。",
                verification="通过故障注入、状态切换和接口测试验证：产生通信故障、中断异常或项目定义诊断事件后，读取接口应返回对应状态；故障清除后应反映更新结果。",
                source=[{"document": "SRS Template", "chunk_id": "DEFAULT-DIAG-READ", "evidence": "Fault/diagnostic handling requires a readable fault or diagnostic status interface."}],
            )
        )

    return result


def dedup_interface_requirements(
    requirements: list[EngineeringRequirement],
    module: str,
) -> tuple[list[EngineeringRequirement], list[ValidationFinding]]:
    """Deduplicate interface requirements by function name.

    Rules:
    1. Interfaces without a ``Gp_xxx_xxx`` function name → validation warning.
    2. Duplicate function names → keep the most complete one (most non-empty
       fields: description, input, output, exception, constraint, verification).
    3. Profile-injected interfaces (semantic_id starts with DEFAULT-) lose
       to datasheet-produced interfaces when both have the same function name.
    """
    findings: list[ValidationFinding] = []
    by_fn: dict[str, list[EngineeringRequirement]] = {}
    for req in requirements:
        if req.requirement_type != "interface":
            continue
        fn = req.function_name.strip()
        if not fn:
            findings.append(ValidationFinding(
                severity="warning", rule="interface-naming",
                rule_group="naming", status="failed",
                requirement_ids=[req.requirement_id],
                message=f"接口需求 {req.requirement_id}（{req.title}）缺少 function_name",
                recommendation="补充 function_name，格式: Gp_{module}_Xxx",
            ))
            continue
        if not re.match(r"^Gp_[A-Za-z0-9]+_[A-Za-z0-9]+$", fn):
            findings.append(ValidationFinding(
                severity="warning", rule="interface-naming",
                rule_group="naming", status="failed",
                requirement_ids=[req.requirement_id],
                message=f"接口 {req.requirement_id}（{req.title}）function_name 格式不符合 Gp_xxx_xxx: {fn}",
                recommendation="修正 function_name 为 Gp_{module}_Xxx 格式",
            ))
        by_fn.setdefault(fn, []).append(req)

    # Score: higher = more complete.  Default-injected gets penalty.
    def _score(r: EngineeringRequirement) -> int:
        s = 0
        for field in (r.description, r.input, r.output, r.exception,
                      r.constraint, r.verification, r.pre_condition, r.trigger):
            if field and field.strip():
                s += 1
        if r.semantic_id.startswith("DEFAULT-"):
            s -= 5  # Heavy penalty — profile/datasheet items preferred
        return s

    remove_ids: set[str] = set()
    for fn, reqs in by_fn.items():
        if len(reqs) <= 1:
            continue
        reqs_sorted = sorted(reqs, key=_score, reverse=True)
        best = reqs_sorted[0]
        for r in reqs_sorted[1:]:
            remove_ids.add(r.requirement_id)
            findings.append(ValidationFinding(
                severity="info", rule="interface-dedup",
                rule_group="consistency", status="failed",
                requirement_ids=[r.requirement_id],
                message=(
                    f"接口 {r.requirement_id}（{r.title}）与 {best.requirement_id}（{best.title}）"
                    f" 共享 function_name {fn}，已去重保留后者"
                ),
                recommendation="无操作；自动去重已处理",
            ))

    if remove_ids:
        requirements = [r for r in requirements if r.requirement_id not in remove_ids]
    return requirements, findings


def _constraints_markdown(findings: list[ValidationFinding]) -> list[str]:
    lines = ["## 10. Constraints", ""]
    if not findings:
        return lines + ["No validation warnings injected.", ""]
    lines += ["| Severity | Rule | Requirement | Finding | Recommendation |", "|---|---|---|---|---|"]
    for finding in findings:
        ids = ", ".join(finding.requirement_ids)
        lines.append(
            f"| {finding.severity} | {finding.rule} | {ids} | {finding.message} | {finding.recommendation} |"
        )
    lines.append("")
    return lines


def _verification_markdown(requirements: list[EngineeringRequirement]) -> list[str]:
    lines = ["## 11. Verification Strategy", "", "| Requirement | Verification |", "|---|---|"]
    for req in requirements:
        lines.append(f"| {req.requirement_id} | {req.verification} |")
    lines.append("")
    return lines


def _traceability_markdown(requirements: list[EngineeringRequirement]) -> list[str]:
    lines = ["## 12. Traceability", "", "### Trace Matrix", "", "| Source | Requirement ID |", "|---|---|"]
    for source, req_id in trace_matrix(requirements):
        lines.append(f"| {source} | {req_id} |")
    lines += ["", "### Coverage Matrix", "", "| Requirement | Source | Validation | Test |", "|---|---|---|---|"]
    for req_id, source, validation, test in coverage_matrix(requirements):
        lines.append(f"| {req_id} | {source} | {validation} | {test} |")
    lines.append("")
    return lines


def _source_summary(req: EngineeringRequirement) -> str:
    if not req.source:
        return ""
    return "; ".join(dict.fromkeys(_public_source_name(source) for source in req.source if _public_source_name(source)))


def _public_source_name(source: dict[str, Any]) -> str:
    document = source.get("document", "")
    chunk_id = source.get("chunk_id", "")
    if document == "RequirementCandidate" or str(chunk_id).startswith("CAND-"):
        return "芯片手册"
    if document == "RequirementPlan":
        return "芯片手册"
    if document == "SRS Template":
        return "需求规范模板"
    return document or chunk_id or "输入材料"


def _markdown_subset_to_html(markdown: str) -> str:
    lines = markdown.splitlines()
    html_lines: list[str] = []
    in_table = False
    for line in lines:
        if line.startswith("|") and line.endswith("|"):
            cells = [escape(cell.strip()) for cell in line.strip("|").split("|")]
            if all(set(cell) <= {"-", ":"} for cell in cells):
                continue
            if not in_table:
                html_lines.append("<table>")
                in_table = True
            tag = "th" if not html_lines[-1].startswith("<tr>") and len(html_lines) >= 1 else "td"
            html_lines.append("<tr>" + "".join(f"<{tag}>{cell}</{tag}>" for cell in cells) + "</tr>")
            continue
        if in_table:
            html_lines.append("</table>")
            in_table = False
        if line.startswith("# "):
            html_lines.append(f"<h1>{escape(line[2:])}</h1>")
        elif line.startswith("## "):
            html_lines.append(f"<h2>{escape(line[3:])}</h2>")
        elif line.startswith("### "):
            html_lines.append(f"<h3>{escape(line[4:])}</h3>")
        elif line.startswith("- "):
            html_lines.append(f"<p>{escape(line[2:])}</p>")
        elif line:
            html_lines.append(f"<p>{escape(line)}</p>")
    if in_table:
        html_lines.append("</table>")
    return "\n".join(html_lines)


def _add_requirement_docx(doc: Any, req: EngineeringRequirement) -> None:
    doc.add_heading(f"{req.requirement_id} {req.title}", level=2)
    rows = [
        ("Requirement ID", req.requirement_id),
        ("Requirement Type", req.requirement_type),
        ("Description", req.description),
        ("Pre-condition", req.pre_condition),
        ("Trigger", req.trigger),
        ("Input", req.input),
        ("Output", req.output),
        ("Exception", req.exception),
        ("Constraint", req.constraint),
        ("Verification", req.verification),
        ("Source", _source_summary(req)),
    ]
    if req.validation:
        rows.append(("Warning", "; ".join(item["message"] for item in req.validation)))
    _add_matrix_docx(doc, ["Field", "Value"], rows)


def _add_findings_table_docx(doc: Any, findings: list[ValidationFinding]) -> None:
    if not findings:
        doc.add_paragraph("No validation warnings injected.")
        return
    rows = [
        [
            finding.severity,
            finding.rule,
            ", ".join(finding.requirement_ids),
            finding.message,
            finding.recommendation,
        ]
        for finding in findings
    ]
    _add_matrix_docx(doc, ["Severity", "Rule", "Requirement", "Finding", "Recommendation"], rows)


def _add_matrix_docx(doc: Any, headers: list[str], rows: list[list[str] | tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
